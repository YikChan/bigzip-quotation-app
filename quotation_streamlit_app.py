# quotation_streamlit_app.py
import streamlit as st
from docx import Document
from docx.shared import Inches
from PIL import Image
import tempfile
import os

st.set_page_config(page_title="Bigzip Quotation Generator", layout="wide")
st.title("🧾 Bigzip Quotation & Confirmation Generator")

# Section A
st.header("Section A: Client Information")
client_name = st.text_input("Client Name")
company_name = st.text_input("Company Name")
company_address = st.text_area("Company Address")
telephone = st.text_input("Telephone")
date = st.date_input("Date")

# Section B & C
st.header("Section B & C: Quotation and Job Info")
quotation_no = st.text_input("Quotation Number")
job_name = st.text_input("Job Name")
job_no = st.text_input("Job Number")

# Section D/E - Work Items
st.header("Section D & E: Work Items")
items = []
with st.form("item_form", clear_on_submit=False):
    desc = st.text_input("Description", key="desc")
    qty = st.number_input("Quantity", min_value=0.0, format="%.2f", key="qty")
    unit_price = st.number_input("Unit Price", min_value=0.0, format="%.2f", key="price")
    submitted = st.form_submit_button("Add Item")
    if submitted:
        st.session_state.setdefault("item_list", []).append({
            "desc": desc,
            "qty": qty,
            "price": unit_price,
            "subtotal": qty * unit_price
        })

if "item_list" not in st.session_state:
    st.session_state["item_list"] = []

if st.session_state["item_list"]:
    st.table(st.session_state["item_list"])
    if st.button("Clear Items"):
        st.session_state["item_list"] = []

total = sum(item['subtotal'] for item in st.session_state["item_list"])

# Section F - Terms
st.header("Section F: Terms")
terms = [
    "1) Quotation is valid for the next 30 days. Jobs must be completed within 6 months, or it may result in an increase in cost.",
    "2) Full payment should be settled within 30 days of invoice issued.",
    "3) Fees and expenses shown are minimum estimates only. Final fees and expenses shall be shown when invoice is rendered.",
    "4) Any additional work or updated not listed in the quotation will be charged accordingly.",
    "5) Crossed cheque should be made payable to 'Bigzip Design Limited', and sent to '7K King Palace Plaza, 55 King Yip Street, Kwun Tong, Kowloon, Hong Kong'.",
    "6) The signature of both parties shall evidence acceptance of these terms.",
    "7) Minimum 30% of the total amount should be charged for cancellation of job after confirmation."
]
for t in terms:
    st.markdown(f"- {t}")

# Section G - Stamp Upload
st.header("Section G: Company Stamp")
stamp_file = st.file_uploader("Upload Stamp Image", type=["png", "jpg", "jpeg"])

# Section H/I - Auto Footer
footer_company = company_name
footer_info = "Bigzip Design Limited\n7K King Palace Plaza, 55 King Yip Street, Kwun Tong, Kowloon, Hong Kong\nTel : 3142 7118 / Fax : 3142 7299 / Email : services@bigzip.com.hk"

# Export Button
if st.button("Generate Quotation (.docx)"):
    doc = Document()
    doc.add_heading("Quotation & Confirmation", 0)
    doc.add_paragraph(f"To: {client_name}")
    doc.add_paragraph(f"Company: {company_name}\n{company_address}")
    doc.add_paragraph(f"Tel: {telephone}")
    doc.add_paragraph(f"Date: {date}")
    doc.add_paragraph(f"Quotation No: {quotation_no}")
    doc.add_paragraph(f"Job Name: {job_name} | Job No: {job_no}")
    doc.add_paragraph("\nWork Details:")

    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Description'
    hdr_cells[1].text = 'Quantity'
    hdr_cells[2].text = 'Unit Price'
    hdr_cells[3].text = 'Subtotal'
    for item in st.session_state["item_list"]:
        row = table.add_row().cells
        row[0].text = item["desc"]
        row[1].text = str(item["qty"])
        row[2].text = f"{item['price']:.2f}"
        row[3].text = f"{item['subtotal']:.2f}"

    doc.add_paragraph(f"\nTotal: HK${total:.2f}")
    doc.add_paragraph("\nTerms:")
    for t in terms:
        doc.add_paragraph(t, style='List Number')

    if stamp_file:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp_img:
            tmp_img.write(stamp_file.read())
            tmp_img_path = tmp_img.name
        doc.add_paragraph("\nStamp:")
        doc.add_picture(tmp_img_path, width=Inches(1.5))

    doc.add_paragraph(f"\n{footer_company}")
    doc.add_paragraph(footer_info)

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_doc:
        doc.save(tmp_doc.name)
        tmp_doc_path = tmp_doc.name

    with open(tmp_doc_path, "rb") as f:
        st.download_button("📥 Download Quotation", f, file_name="quotation.docx")
